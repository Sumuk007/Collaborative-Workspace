from sqlalchemy.orm import Session, joinedload
from app.models.document import Document
from app.models.user import User
from app.models.document_collaborator import DocumentCollaborator
from app.models.share_link import ShareLink
from app.schemas.document import DocumentCreate, DocumentUpdate
from typing import List, Optional, Literal
from datetime import datetime, timedelta
import secrets
from io import BytesIO
from bs4 import BeautifulSoup
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from html import unescape


def create_document(db: Session, document_data: DocumentCreate, owner_id: int) -> Document:
    """Create a new document"""
    # Sanitize title
    title = document_data.title.strip()
    
    if not title:
        raise ValueError("Title cannot be empty")
    
    # Check for duplicate title for this user
    existing = db.query(Document).filter(
        Document.owner_id == owner_id,
        Document.title == title
    ).first()
    
    if existing:
        raise ValueError(f"Document with title '{title}' already exists")
    
    db_document = Document(
        title=title,
        content=document_data.content or "",
        content_type=document_data.content_type or "plain",
        content_blocks=document_data.content_blocks,
        styles=document_data.styles,
        owner_id=owner_id
    )
    db.add(db_document)
    db.flush()  # Get the document ID before creating collaborator
    
    # Add owner as collaborator with 'owner' role
    owner_collab = DocumentCollaborator(
        document_id=db_document.id,
        user_id=owner_id,
        role="owner"
    )
    db.add(owner_collab)
    db.commit()
    db.refresh(db_document)
    return db_document


def get_document_by_id(db: Session, document_id: int) -> Optional[Document]:
    """Get a single document by ID"""
    return db.query(Document).filter(Document.id == document_id).first()


def get_user_documents(db: Session, user_id: int, skip: int = 0, limit: int = 100) -> List[Document]:
    """Get all documents accessible to a user (owned or shared)"""
    return db.query(Document).join(DocumentCollaborator).filter(
        DocumentCollaborator.user_id == user_id
    ).order_by(Document.updated_at.desc()).offset(skip).limit(limit).all()


def get_all_documents(db: Session, skip: int = 0, limit: int = 100) -> List[Document]:
    """Get all documents (for admin or public view)"""
    return db.query(Document).offset(skip).limit(limit).all()


def update_document(db: Session, document_id: int, document_data: DocumentUpdate) -> Optional[Document]:
    """Update a document"""
    db_document = db.query(Document).filter(Document.id == document_id).first()
    
    if not db_document:
        return None
    
    # Update only provided fields
    update_data = document_data.model_dump(exclude_unset=True)
    
    # Validate title if being updated
    if 'title' in update_data:
        title = update_data['title'].strip() if update_data['title'] else None
        
        if not title:
            raise ValueError("Title cannot be empty")
        
        # Check for duplicate title for this user (excluding current document)
        existing = db.query(Document).filter(
            Document.owner_id == db_document.owner_id,
            Document.title == title,
            Document.id != document_id
        ).first()
        
        if existing:
            raise ValueError(f"Document with title '{title}' already exists")
        
        update_data['title'] = title
    
    for field, value in update_data.items():
        setattr(db_document, field, value)
    
    db.commit()
    db.refresh(db_document)
    return db_document


def delete_document(db: Session, document_id: int) -> bool:
    """Delete a document"""
    db_document = db.query(Document).filter(Document.id == document_id).first()
    
    if not db_document:
        return False
    
    db.delete(db_document)
    db.commit()
    return True


def search_documents(db: Session, query: str, user_id: Optional[int] = None, skip: int = 0, limit: int = 100) -> List[Document]:
    """Search documents by title or content"""
    search_query = db.query(Document).filter(
        (Document.title.ilike(f"%{query}%")) | (Document.content.ilike(f"%{query}%"))
    )
    
    if user_id:
        search_query = search_query.filter(Document.owner_id == user_id)
    
    return search_query.offset(skip).limit(limit).all()


def get_user_role_for_document(db: Session, document_id: int, user_id: int) -> Optional[str]:
    """Get user's role for a specific document"""
    collaborator = db.query(DocumentCollaborator).filter(
        DocumentCollaborator.document_id == document_id,
        DocumentCollaborator.user_id == user_id
    ).first()
    
    return collaborator.role if collaborator else None


def add_collaborator(db: Session, document_id: int, user_id: int, role: Literal["editor", "reader"]) -> DocumentCollaborator:
    """Add a collaborator to a document"""
    # Check if collaborator already exists
    existing = db.query(DocumentCollaborator).filter(
        DocumentCollaborator.document_id == document_id,
        DocumentCollaborator.user_id == user_id
    ).first()
    
    if existing:
        raise ValueError(f"User already has '{existing.role}' role on this document")
    
    # Validate role
    if role not in ["editor", "reader"]:
        raise ValueError("Role must be 'editor' or 'reader'")
    
    collaborator = DocumentCollaborator(
        document_id=document_id,
        user_id=user_id,
        role=role
    )
    db.add(collaborator)
    db.commit()
    db.refresh(collaborator)
    return collaborator


def remove_collaborator(db: Session, document_id: int, user_id: int) -> bool:
    """Remove a collaborator from a document"""
    collaborator = db.query(DocumentCollaborator).filter(
        DocumentCollaborator.document_id == document_id,
        DocumentCollaborator.user_id == user_id,
        DocumentCollaborator.role != "owner"  # Cannot remove owner
    ).first()
    
    if not collaborator:
        return False
    
    db.delete(collaborator)
    db.commit()
    return True


def get_document_collaborators(db: Session, document_id: int) -> List[DocumentCollaborator]:
    """Get all collaborators for a document with user information"""
    return db.query(DocumentCollaborator).options(
        joinedload(DocumentCollaborator.user)
    ).filter(
        DocumentCollaborator.document_id == document_id
    ).all()


def create_share_link(db: Session, document_id: int, role: str, created_by: int, expires_in_hours: Optional[int] = None) -> ShareLink:
    """Create a shareable link for a document"""
    token = secrets.token_urlsafe(32)
    
    expires_at = None
    if expires_in_hours:
        expires_at = datetime.utcnow() + timedelta(hours=expires_in_hours)
    
    share_link = ShareLink(
        document_id=document_id,
        token=token,
        role=role,
        created_by=created_by,
        expires_at=expires_at
    )
    db.add(share_link)
    db.commit()
    db.refresh(share_link)
    return share_link


def get_share_link_by_token(db: Session, token: str) -> Optional[ShareLink]:
    """Get a share link by token"""
    share_link = db.query(ShareLink).filter(
        ShareLink.token == token,
        ShareLink.is_active == 1
    ).first()
    
    if not share_link:
        return None
    
    # Check if expired
    if share_link.expires_at and share_link.expires_at < datetime.utcnow():
        return None
    
    return share_link


def accept_share_link(db: Session, token: str, user_id: int) -> Optional[DocumentCollaborator]:
    """Accept a share link and add user as collaborator or update their role"""
    share_link = get_share_link_by_token(db, token)
    
    if not share_link:
        return None
    
    # Check if user is already a collaborator
    existing = db.query(DocumentCollaborator).filter(
        DocumentCollaborator.document_id == share_link.document_id,
        DocumentCollaborator.user_id == user_id
    ).first()
    
    if existing:
        # Don't allow changing owner role
        if existing.role == "owner":
            raise ValueError("You are already the owner of this document")
        
        # If the new role is different, update it
        if existing.role != share_link.role:
            existing.role = share_link.role
            db.commit()
            db.refresh(existing)
            return existing
        else:
            # Same role, just return the existing collaborator
            raise ValueError(f"You are already a {existing.role} on this document")
    
    # Add user as collaborator
    collaborator = DocumentCollaborator(
        document_id=share_link.document_id,
        user_id=user_id,
        role=share_link.role
    )
    db.add(collaborator)
    db.commit()
    db.refresh(collaborator)
    return collaborator


def revoke_share_link(db: Session, document_id: int, token: str) -> bool:
    """Revoke/disable a share link"""
    share_link = db.query(ShareLink).filter(
        ShareLink.document_id == document_id,
        ShareLink.token == token
    ).first()
    
    if not share_link:
        return False
    
    share_link.is_active = 0
    db.commit()
    return True


def update_collaborator_role(db: Session, document_id: int, user_id: int, new_role: Literal["editor", "reader"]) -> Optional[DocumentCollaborator]:
    """Update a collaborator's role - cannot change owner role"""
    collaborator = db.query(DocumentCollaborator).filter(
        DocumentCollaborator.document_id == document_id,
        DocumentCollaborator.user_id == user_id
    ).first()
    
    if not collaborator:
        return None
    
    # Cannot change owner role
    if collaborator.role == "owner":
        raise ValueError("Cannot change owner's role")
    
    # Validate new role
    if new_role not in ["editor", "reader"]:
        raise ValueError("Role must be 'editor' or 'reader'")
    
    collaborator.role = new_role
    db.commit()
    db.refresh(collaborator)
    return collaborator


def get_document_share_links(db: Session, document_id: int) -> List[ShareLink]:
    """Get all share links for a document"""
    return db.query(ShareLink).filter(
        ShareLink.document_id == document_id,
        ShareLink.is_active == 1
    ).all()


def parse_html_content(html_content: str) -> List[dict]:
    """Parse HTML content and extract text with formatting"""
    if not html_content:
        return []
        
    try:
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        elements = []
        seen_texts_with_context = set()  # Track text with type to allow duplicates in different contexts
        
        def process_element(element, is_list_item=False):
            """Recursively process an element"""
            # Handle direct text nodes
            if isinstance(element, str):
                text = element.strip()
                if text:
                    context_key = f"text_{text}"
                    if context_key not in seen_texts_with_context:
                        seen_texts_with_context.add(context_key)
                        elements.append({
                            'type': 'li' if is_list_item else 'p',
                            'text': unescape(text),
                            'align': 'left',
                            'bold': False,
                            'italic': False,
                            'underline': False,
                            'is_list_item': is_list_item
                        })
                return
            
            # Handle tags
            if not hasattr(element, 'name'):
                return
            
            if element.name == 'br':
                elements.append({'type': 'break', 'text': ''})
                return
            
            # Handle lists specially - process each li separately
            if element.name in ['ul', 'ol']:
                list_items = element.find_all('li', recursive=False)
                for idx, li in enumerate(list_items, 1):
                    text = li.get_text(strip=True)
                    if not text:
                        continue
                    
                    # Create context key to allow same text in different contexts
                    context_key = f"li_{element.name}_{text}"
                    if context_key in seen_texts_with_context:
                        continue
                    
                    seen_texts_with_context.add(context_key)
                    
                    # Check for formatting
                    style = li.get('style', '') or ''
                    is_bold = bool(li.find(['strong', 'b'])) or 'font-weight: bold' in style or 'font-weight:bold' in style
                    is_italic = bool(li.find(['em', 'i'])) or 'font-style: italic' in style or 'font-style:italic' in style
                    # Check if text is underlined - check if wrapped in <u> or has style
                    u_elem = li.find('u')
                    is_underline = (li.name == 'u' or 
                                  (hasattr(li.parent, 'name') and li.parent.name == 'u') or 
                                  (u_elem and text in u_elem.get_text(strip=True)) or
                                  'text-decoration: underline' in style or 
                                  'text-decoration:underline' in style)
                    
                    # Extract font size for list items
                    import re
                    font_size = None
                    for descendant in li.descendants:
                        if hasattr(descendant, 'get'):
                            desc_style = descendant.get('style', '')
                            if desc_style and ('font-size:' in desc_style or 'font-size :' in desc_style):
                                match = re.search(r'font-size\s*:\s*(\d+)px', desc_style)
                                if match:
                                    font_size = int(match.group(1))
                                    break
                    
                    list_elem = {
                        'type': 'li',
                        'text': unescape(text),
                        'align': 'left',
                        'bold': is_bold,
                        'italic': is_italic,
                        'underline': is_underline,
                        'is_list_item': True,
                        'list_type': element.name,  # 'ul' or 'ol'
                        'list_index': idx  # For ordered lists
                    }
                    if font_size:
                        list_elem['font_size'] = font_size
                    
                    elements.append(list_elem)
                return
            
            # Check if this element contains lists - process children separately
            child_lists = element.find_all(['ul', 'ol'], recursive=False)
            if child_lists:
                # Process children in order
                for child in element.children:
                    process_element(child)
                return
            
            # Get text content for other elements (no lists inside)
            text = element.get_text(strip=True)
            if not text:
                return
            
            # Skip underline tags that only contain breaks or whitespace
            if element.name == 'u' and (not text or text.isspace()):
                return
            
            # Create context key to allow same text in different contexts
            context_key = f"{element.name}_{text}"
            if context_key in seen_texts_with_context:
                return
            
            seen_texts_with_context.add(context_key)
            
            # Determine alignment from style
            style = element.get('style', '') or ''
            align = 'left'
            if 'text-align: center' in style or 'text-align:center' in style:
                align = 'center'
            elif 'text-align: right' in style or 'text-align:right' in style:
                align = 'right'
            elif 'text-align: justify' in style or 'text-align:justify' in style:
                align = 'justify'
            
            # Check for formatting
            is_bold = bool(element.find(['strong', 'b'])) or 'font-weight: bold' in style or 'font-weight:bold' in style
            is_italic = bool(element.find(['em', 'i'])) or 'font-style: italic' in style or 'font-style:italic' in style
            # Check if text is underlined - check if wrapped in <u> or has style
            u_elem = element.find('u')
            is_underline = (element.name == 'u' or 
                          (hasattr(element.parent, 'name') and element.parent.name == 'u') or 
                          (u_elem and text in u_elem.get_text(strip=True)) or
                          'text-decoration: underline' in style or 
                          'text-decoration:underline' in style)
            
            # Extract font size - check element itself and all descendants
            import re
            font_size = None
            
            # Check current element's style
            if 'font-size:' in style or 'font-size :' in style:
                match = re.search(r'font-size\s*:\s*(\d+)px', style)
                if match:
                    font_size = int(match.group(1))
            
            # Check all descendants for font-size
            if not font_size:
                for descendant in element.descendants:
                    if hasattr(descendant, 'get'):
                        desc_style = descendant.get('style', '')
                        if desc_style and ('font-size:' in desc_style or 'font-size :' in desc_style):
                            match = re.search(r'font-size\s*:\s*(\d+)px', desc_style)
                            if match:
                                font_size = int(match.group(1))
                                break
            
            elem_type = 'p' if element.name in ['div', 'span'] else element.name
            
            elem_data = {
                'type': elem_type,
                'text': unescape(text),
                'align': align,
                'bold': is_bold,
                'italic': is_italic,
                'underline': is_underline
            }
            if font_size:
                elem_data['font_size'] = font_size
            
            elements.append(elem_data)
        
        # Process all children of the root
        for child in soup.children:
            process_element(child)
        
        return elements
    except Exception as e:
        # Return at least something to avoid complete failure
        return [{'type': 'p', 'text': 'Error parsing document content', 'align': 'left', 'bold': False, 'italic': False, 'underline': False}]


def export_document_to_pdf(db: Session, document_id: int, user_id: int) -> BytesIO:
    """
    Export document to PDF format
    
    Args:
        db: Database session
        document_id: ID of the document to export
        user_id: ID of the user requesting export
        
    Returns:
        BytesIO: PDF file buffer
        
    Raises:
        ValueError: If document not found or access denied
    """
    # Get document
    document = get_document_by_id(db, document_id)
    if not document:
        raise ValueError("Document not found")
    
    # Check if user has access
    user_role = get_user_role_for_document(db, document_id, user_id)
    if not user_role:
        raise ValueError("Access denied")
    
    # Create PDF buffer
    buffer = BytesIO()
    
    # Create PDF document
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72
    )
    
    # Build story (content)
    story = []
    styles = getSampleStyleSheet()
    
    # Add custom styles
    styles.add(ParagraphStyle(
        name='CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
        alignment=TA_CENTER
    ))
    
    styles.add(ParagraphStyle(
        name='CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        spaceAfter=12,
        spaceBefore=12
    ))
    
    styles.add(ParagraphStyle(name='CustomBody', parent=styles['BodyText'], fontSize=11, alignment=TA_LEFT, spaceAfter=12))
    styles.add(ParagraphStyle(name='CustomBodyCenter', parent=styles['BodyText'], fontSize=11, alignment=TA_CENTER, spaceAfter=12))
    styles.add(ParagraphStyle(name='CustomBodyRight', parent=styles['BodyText'], fontSize=11, alignment=TA_RIGHT, spaceAfter=12))
    styles.add(ParagraphStyle(name='CustomBodyJustify', parent=styles['BodyText'], fontSize=11, alignment=TA_JUSTIFY, spaceAfter=12))
    
    styles.add(ParagraphStyle(
        name='Blockquote',
        parent=styles['BodyText'],
        fontSize=11,
        leftIndent=20,
        rightIndent=20,
        textColor=colors.HexColor('#666666'),
        spaceAfter=12
    ))
    
    # Parse and add content
    if document.content:
        elements = parse_html_content(document.content)
        
        for elem in elements:
            if elem['type'] == 'break':
                story.append(Spacer(1, 0.1 * inch))
                continue
            
            text = elem['text']
            
            # Add bullet for list items
            if elem.get('is_list_item'):
                if elem.get('list_type') == 'ol':
                    text = f"{elem.get('list_index', 1)}. {text}"
                else:
                    text = f"• {text}"
            
            # Apply inline formatting
            if elem.get('bold'):
                text = f"<b>{text}</b>"
            if elem.get('italic'):
                text = f"<i>{text}</i>"
            if elem.get('underline'):
                text = f"<u>{text}</u>"
            
            # Choose style based on element type and alignment
            if elem['type'] in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                style = styles['CustomHeading']
            elif elem['type'] == 'blockquote':
                style = styles['Blockquote']
            else:
                if elem['align'] == 'center':
                    style = styles['CustomBodyCenter']
                elif elem['align'] == 'right':
                    style = styles['CustomBodyRight']
                elif elem['align'] == 'justify':
                    style = styles['CustomBodyJustify']
                else:
                    style = styles['CustomBody']
                
                # Create custom style with font size if specified (including list items)
                if elem.get('font_size'):
                    font_size = elem.get('font_size')
                    style_name = f"Custom_{font_size}_{elem['align']}_{elem.get('type', 'p')}"
                    custom_style = ParagraphStyle(
                        name=style_name,
                        parent=style,
                        fontSize=font_size,
                        leading=font_size * 1.2,  # Line height = 120% of font size
                        spaceAfter=font_size * 0.5  # Space after = 50% of font size
                    )
                    style = custom_style
            
            story.append(Paragraph(text, style))
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    
    return buffer


def export_document_to_docx(db: Session, document_id: int, user_id: int) -> BytesIO:
    """
    Export document to Word (DOCX) format
    
    Args:
        db: Database session
        document_id: ID of the document to export
        user_id: ID of the user requesting export
        
    Returns:
        BytesIO: DOCX file buffer
        
    Raises:
        ValueError: If document not found or access denied
    """
    # Get document
    document = get_document_by_id(db, document_id)
    if not document:
        raise ValueError("Document not found")
    
    # Check if user has access
    user_role = get_user_role_for_document(db, document_id, user_id)
    if not user_role:
        raise ValueError("Access denied")
    
    # Create Word document
    doc = DocxDocument()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Parse and add content
    if document.content:
        elements = parse_html_content(document.content)
        
        for elem in elements:
            if elem['type'] == 'break':
                doc.add_paragraph()
                continue
            
            # Determine paragraph type
            if elem['type'] in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(elem['type'][1])
                para = doc.add_heading(elem['text'], level=level)
            elif elem.get('is_list_item'):
                # Add as bullet or numbered point
                list_style = 'List Number' if elem.get('list_type') == 'ol' else 'List Bullet'
                para = doc.add_paragraph(elem['text'], style=list_style)
                run = para.runs[0] if para.runs else para.add_run(elem['text'])
                
                # Apply formatting to the run
                if elem.get('bold'):
                    run.bold = True
                if elem.get('italic'):
                    run.italic = True
                if elem.get('underline'):
                    run.underline = True
                
                # Apply font size
                if elem.get('font_size'):
                    run.font.size = Pt(elem['font_size'])
                else:
                    run.font.size = Pt(11)
            else:
                para = doc.add_paragraph()
                run = para.add_run(elem['text'])
                
                # Apply formatting
                if elem.get('bold'):
                    run.bold = True
                if elem.get('italic'):
                    run.italic = True
                if elem.get('underline'):
                    run.underline = True
                
                # Apply font size
                if elem.get('font_size'):
                    run.font.size = Pt(elem['font_size'])
                else:
                    run.font.size = Pt(11)
                
                # Apply blockquote styling
                if elem['type'] == 'blockquote':
                    para.paragraph_format.left_indent = Inches(0.5)
                    para.paragraph_format.right_indent = Inches(0.5)
                    run.font.color.rgb = RGBColor(102, 102, 102)
            
            # Set alignment
            if elem['align'] == 'center':
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif elem['align'] == 'right':
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif elem['align'] == 'justify':
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            else:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer
